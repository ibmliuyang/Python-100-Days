import os

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from openai import OpenAI
from concurrent.futures import ThreadPoolExecutor, as_completed
import time

# 从环境变量中获取 API 密钥
api_key = "sk-d0a0e5d5358e4d579ac8c5afc2961c37"

# 初始化 DeepSeek 客户端
client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")

# 定义翻译函数（带重试机制）
def translate_text(text, retries=3):
    for attempt in range(retries):
        try:
            response = client.chat.completions.create(
                model="deepseek-chat",
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "你好，现在你是一名人类学、社会学方面的专家，擅长将人类学和社会学方面的图书从英文译成中文。"
                            "今天我们翻译人类学学者 Tim Ingold 和 Jo Lee Vergunst 编辑的作品《Ways of Walking: Ethnography and Practice on Foot》，"
                            "中文名字是《行走之道：步行中的民族志与实践》。具体规则如下：\n"
                            "- 翻译时要准确传达原文的事实和背景。\n"
                            "- 即使上意译也要保留原始段落格式，以及保留人类学或者社会学术语。保留公司和博物馆缩写等。\n"
                            "- 人名要翻译，中文人名后加括号，括号里写清楚英文原名。\n"
                            "- 同时要保留引用的论文，例如 [20] 这样的引用。\n"
                            "- 在翻译专业术语时，第一次出现时要在括号里面写上英文原文，之后就可以只写中文了。\n"
                            "- 除术语和专业表达以外，其余表述要求简单、清楚，不要过多有文学色彩的表达。"
                            "策略：分四步进行翻译工作，并打印每步的结果：\n"
                            "1. 根据英文内容直译，保持原有格式，不要遗漏任何信息。\n"
                            "2. 根据第一步直译的结果，指出其中存在的具体问题，要准确描述，不宜笼统的表示，也不需要增加原文不存在的内容或格式，包括不仅限于：\n"
                            "   - 不符合中文表达习惯，明确指出不符合的地方。\n"
                            "   - 语句不通顺，指出位置，不需要给出修改意见，意译时修复。\n"
                            "   - 晦涩难懂，不易理解，可以尝试给出解释。\n"
                            "3. 根据第一步直译的结果和第二步指出的问题，重新进行意译，保证内容的原意的基础上，使其更易于理解，更符合中文的表达习惯，同时保持原有的格式不变。\n"
                            "4. 只保留第3部的内容即可，1、2不需要输出 。\n"
                        ),
                    },
                    {"role": "user", "content": text},
                ],
                stream=False,
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"API 调用失败，重试 {attempt + 1}/{retries}：{e}")
            time.sleep(2)  # 等待 2 秒后重试
    return None  # 重试多次后仍失败，返回 None

# 动态分块函数（按段落分块，保留上下文）
def split_text_into_chunks(text, max_chunk_size=10000, overlap_size=200, is_english=False):
    paragraphs = text.split("\n")
    chunks = []
    current_chunk = ""
    overlap_text = ""

    # 根据语言类型调整 token 比例
    token_ratio = 0.3 if is_english else 0.6

    for paragraph in paragraphs:
        # 计算当前 chunk 的 token 数量
        current_chunk_tokens = len(current_chunk) * token_ratio
        paragraph_tokens = len(paragraph) * token_ratio

        # 如果当前 chunk 加上新段落不超过最大大小，则添加到当前 chunk
        if current_chunk_tokens + paragraph_tokens <= max_chunk_size:
            current_chunk += paragraph + "\n"
        else:
            # 否则保存当前 chunk，并开始新的 chunk
            if current_chunk:
                chunks.append(overlap_text + current_chunk)
            # 保留最后一部分作为上下文
            overlap_text = current_chunk[-overlap_size:] if len(current_chunk) > overlap_size else current_chunk
            current_chunk = paragraph + "\n"

    # 添加最后一个 chunk
    if current_chunk:
        chunks.append(overlap_text + current_chunk)

    return chunks

# 并发翻译函数（带进度打印）
def translate_concurrently(text, max_chunk_size=10000, max_workers=10, is_english=False):
    # 动态分块
    chunks = split_text_into_chunks(text, max_chunk_size, is_english=is_english)
    translated_chunks = [None] * len(chunks)  # 用于存储翻译结果，保持顺序
    total_tasks = len(chunks)  # 总任务数
    completed_tasks = 0  # 已完成任务数

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        # 提交任务，并记录每个任务的索引
        future_to_index = {executor.submit(translate_text, chunk): i for i, chunk in enumerate(chunks)}

        # 按任务完成顺序获取结果
        for future in as_completed(future_to_index):
            index = future_to_index[future]  # 获取任务的原始索引
            translated_chunks[index] = future.result()  # 将结果放入正确的位置
            completed_tasks += 1  # 更新已完成任务数
            # 打印进度
            print(f"翻译进度: {completed_tasks}/{total_tasks} ({(completed_tasks / total_tasks) * 100:.2f}%)")

    # 按原始顺序拼接翻译结果
    translated_text = "".join(translated_chunks)
    return translated_text

# 读取 Word 文档内容
def read_word_file(file_path):
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

# 将翻译结果写入 Word 文档（设置字体为仿宋）
def write_to_word_file(text, output_path):
    doc = Document()
    for paragraph in text.split("\n"):
        p = doc.add_paragraph(paragraph)
        # 设置字体为仿宋
        for run in p.runs:
            run.font.name = "仿宋"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")  # 设置中文字体
            run.font.size = Pt(12)  # 设置字号为 12
    doc.save(output_path)

# 主函数
def translate_word_file(input_path, output_path, is_english=False):
    # 读取原始文档
    print("正在读取原始文档...")
    original_text = read_word_file(input_path)
    print("原始文档读取完成。")

    # 并发翻译
    print("正在调用 DeepSeek API 进行翻译...")
    translated_text = translate_concurrently(original_text, is_english=is_english)
    print("翻译完成。")

    # 将翻译结果写入新文档
    print("正在将翻译结果写入新文档...")
    if translated_text is not None:
        print(f"translated_text===:==={translated_text}")
    else:
        print("translated_text is None")
    write_to_word_file(translated_text, output_path)
    print(f"翻译结果已保存到 {output_path}")

# 示例调用
if __name__ == "__main__":
    # 打印当前工作目录
    print(f"当前工作目录: {os.getcwd()}")
    input_file = "w/docx/en999_1.docx"  # 输入文件路径
    output_file = "w/docx/cn991_1.docx"  # 输出文件路径
    is_english = True  # 如果输入文件是英文，设置为 True；如果是中文，设置为 False
    translate_word_file(input_file, output_file, is_english)