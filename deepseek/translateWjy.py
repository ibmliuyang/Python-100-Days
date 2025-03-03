from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from openai import OpenAI
from concurrent.futures import ThreadPoolExecutor, as_completed

# 从环境变量中获取 API 密钥
api_key = "sk-d0a0e5d5358e4d579ac8c5afc2961c37"

# 初始化 DeepSeek 客户端
client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")


# 定义翻译函数
def translate_text(text):
    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[
            {
                "role": "system",
                "content": (
                    "你好，现在你是一名人类学、社会学方面的专家，擅长将人类学和社会学方面的图书从英文译成中文。"
                    "今天我们翻译人类学学者 Tim Ingold 和 Jo Lee Vergunst 编辑的作品《Ways of Walking: Ethnography and Practice on Foot》，"
                    "中文名字是《行走之道：步行中的民族志与实践》。具体规则如下：\n"
                    "- 翻译时要准确传达原文的事实和背景。\n"
                    "- 即使上意译也要保留原始段落格式，以及保留人类学或者社会学术语。保留公司和博物馆缩写等。\n"
                    "- 人名要翻译，中文人名后加括号，括号里写清楚英文原名。\n"
                    "- 同时要保留引用的论文，例如 [20] 这样的引用。\n"
                    "- 在翻译专业术语时，第一次出现时要在括号里面写上英文原文，之后就可以只写中文了。\n"
                    "- 除术语和专业表达以外，其余表述要求简单、清楚，不要过多有文学色彩的表达。"
                    "策略：分三步进行翻译工作，并打印每步的结果：\n"
                    "1. 根据英文内容直译，保持原有格式，不要遗漏任何信息。\n"
                    "2. 根据第一步直译的结果，指出其中存在的具体问题，要准确描述，不宜笼统的表示，也不需要增加原文不存在的内容或格式，包括不仅限于：\n"
                    "   - 不符合中文表达习惯，明确指出不符合的地方。\n"
                    "   - 语句不通顺，指出位置，不需要给出修改意见，意译时修复。\n"
                    "   - 晦涩难懂，不易理解，可以尝试给出解释。\n"
                    "3. 根据第一步直译的结果和第二步指出的问题，重新进行意译，保证内容的原意的基础上，使其更易于理解，更符合中文的表达习惯，同时保持原有的格式不变。\n"
                    " 4. 只保留第3部的内容即可，1、2不需要输出 。\n"

                ),
            },
            {"role": "user", "content": text},
        ],
        stream=False,
    )
    return response.choices[0].message.content


# 并发翻译函数（带进度打印）
def translate_concurrently(text, chunk_size=500, max_workers=10):
    # 将文本分块
    chunks = [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]
    translated_chunks = [None] * len(chunks)  # 用于存储翻译结果，保持顺序
    total_tasks = len(chunks)  # 总任务数
    completed_tasks = 0  # 已完成任务数

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        # 提交任务，并记录每个任务的索引
        future_to_index = {executor.submit(translate_text, chunk): i for i, chunk in enumerate(chunks)}

        # 按任务完成顺序获取结果
        for future in as_completed(future_to_index):
            index = future_to_index[future]  # 获取任务的原始索引
            translated_chunks[index] = future.result()  # 将结果放入正确的位置
            completed_tasks += 1  # 更新已完成任务数
            # 打印进度
            print(f"翻译进度: {completed_tasks}/{total_tasks} ({(completed_tasks / total_tasks) * 100:.2f}%)")

    # 按原始顺序拼接翻译结果
    translated_text = "".join(translated_chunks)
    return translated_text


# 读取 Word 文档内容
def read_word_file(file_path):
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text


# 将翻译结果写入 Word 文档（设置字体为仿宋）
def write_to_word_file(text, output_path):
    doc = Document()
    for paragraph in text.split("\n"):
        p = doc.add_paragraph(paragraph)
        # 设置字体为仿宋
        for run in p.runs:
            run.font.name = "仿宋"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")  # 设置中文字体
            run.font.size = Pt(12)  # 设置字号为 12
    doc.save(output_path)


# 主函数
def translate_word_file(input_path, output_path):
    # 读取原始文档
    print("正在读取原始文档...")
    original_text = read_word_file(input_path)
    print("原始文档读取完成。")

    # 并发翻译
    print("正在调用 DeepSeek API 进行翻译...")
    translated_text = translate_concurrently(original_text)
    print("翻译完成。")

    # 将翻译结果写入新文档
    print("正在将翻译结果写入新文档...")
    write_to_word_file(translated_text, output_path)
    print(f"翻译结果已保存到 {output_path}")


# 示例调用
if __name__ == "__main__":
    input_file = "/Users/ly/Downloads/wjy/副本en.docx"  # 输入文件路径
    output_file = "/Users/ly/Downloads/wjy/cn4.docx"  # 输出文件路径
    translate_word_file(input_file, output_file)