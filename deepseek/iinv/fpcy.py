import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from openai import OpenAI
from concurrent.futures import ThreadPoolExecutor, as_completed
import time
from PyPDF2 import PdfReader  # 新增 PDF 文件支持

# 从环境变量中获取 API 密钥
api_key = "sk-d0a0e5d5358e4d579ac8c5afc2961c37"

# 初始化 DeepSeek 客户端
client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")

# 定义编码函数（带重试机制）
def translate_text(text, retries=3):
    for attempt in range(retries):
        try:
            response = client.chat.completions.create(
                model="deepseek-chat",
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "你是一个资深的java开发工程师，当你阅读完所有的文档，帮我完成'接口列表 （一）接口清单 发票查验（单张发票实时查询下载接口） '的开发工作要求健壮、可扩展、定义合适的枚举、常量。提供一个合适的方法用来完成发票查验，封装合适的入参和出参。"
                        ),
                    },
                    {"role": "user", "content": text},
                ],
                stream=False,
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"API 调用失败，重试 {attempt + 1}/{retries}：{e}")
            time.sleep(2)  # 等待 2 秒后重试
    return None  # 重试多次后仍失败，返回 None

# 动态分块函数（按段落分块，保留上下文）
def split_text_into_chunks(text, max_chunk_size=10000, overlap_size=200, is_english=False):
    paragraphs = text.split("\n")
    chunks = []
    current_chunk = ""
    overlap_text = ""

    # 根据语言类型调整 token 比例
    token_ratio = 0.3 if is_english else 0.6

    for paragraph in paragraphs:
        # 计算当前 chunk 的 token 数量
        current_chunk_tokens = len(current_chunk) * token_ratio
        paragraph_tokens = len(paragraph) * token_ratio

        # 如果当前 chunk 加上新段落不超过最大大小，则添加到当前 chunk
        if current_chunk_tokens + paragraph_tokens <= max_chunk_size:
            current_chunk += paragraph + "\n"
        else:
            # 否则保存当前 chunk，并开始新的 chunk
            if current_chunk:
                chunks.append(overlap_text + current_chunk)
            # 保留最后一部分作为上下文
            overlap_text = current_chunk[-overlap_size:] if len(current_chunk) > overlap_size else current_chunk
            current_chunk = paragraph + "\n"

    # 添加最后一个 chunk
    if current_chunk:
        chunks.append(overlap_text + current_chunk)

    return chunks

# 并发编码函数（带进度打印）
def translate_concurrently(text, max_chunk_size=10000, max_workers=10, is_english=False):
    # 动态分块
    chunks = split_text_into_chunks(text, max_chunk_size, is_english=is_english)
    code_chunks = [None] * len(chunks)  # 用于存储结果，保持顺序
    total_tasks = len(chunks)  # 总任务数
    completed_tasks = 0  # 已完成任务数

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        # 提交任务，并记录每个任务的索引
        future_to_index = {executor.submit(translate_text, chunk): i for i, chunk in enumerate(chunks)}

        # 按任务完成顺序获取结果
        for future in as_completed(future_to_index):
            index = future_to_index[future]  # 获取任务的原始索引
            code_chunks[index] = future.result()  # 将结果放入正确的位置
            completed_tasks += 1  # 更新已完成任务数
            # 打印进度
            print(f"编码进度: {completed_tasks}/{total_tasks} ({(completed_tasks / total_tasks) * 100:.2f}%)")

    # 按原始顺序拼接编码结果
    code_text = "".join(code_chunks)
    return code_text

# 读取 Word 文档内容
def read_word_file(file_path):
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

# 读取 PDF 文档内容
def read_pdf_file(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text

# 将编码结果写入 Word 文档（设置字体为仿宋）
def write_to_word_file(text, output_path):
    doc = Document()
    for paragraph in text.split("\n"):
        p = doc.add_paragraph(paragraph)
        # 设置字体为仿宋
        for run in p.runs:
            run.font.name = "仿宋"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")  # 设置中文字体
            run.font.size = Pt(12)  # 设置字号为 12
    doc.save(output_path)

# 主函数
def code_word_file(input_path, output_path, is_english=False):
    # 检查文件是否存在
    if not os.path.exists(input_path):
        print(f"错误：文件 '{input_path}' 不存在。")
        return

    # 读取原始文档
    print("正在读取原始文档...")
    try:
        if input_path.endswith(".docx"):
            original_text = read_word_file(input_path)
        elif input_path.endswith(".pdf"):
            original_text = read_pdf_file(input_path)
        else:
            print("错误：不支持的文件格式。仅支持 .docx 和 .pdf 文件。")
            return
    except Exception as e:
        print(f"读取文件失败：{e}")
        return
    print("原始文档读取完成。")

    # 并发编码
    print("正在调用 DeepSeek API 进行编码...")
    code_text = translate_concurrently(original_text, is_english=is_english)
    print("编码完成。")

    # 将结果写入新文档
    print("正在将结果写入新文档...")
    if code_text is not None:
        print(f"code_text===:==={code_text}")
        write_to_word_file(code_text, output_path)
        print(f"结果已保存到 {output_path}")
    else:
        print("错误：编码结果为空。")

# 示例调用
if __name__ == "__main__":
    # 打印当前工作目录
    print(f"当前工作目录: {os.getcwd()}")
    input_file = "/Users/ly/ynworkspace/Python-100-Days/deepseek/iinv/乐企发票查验能力说明文档.pdf"  # 输入文件路径
    output_file = "fpcy_kaifa.docx"  # 输出文件路径
    is_english = False  # 如果输入文件是英文，设置为 True；如果是中文，设置为 False
    code_word_file(input_file, output_file, is_english)